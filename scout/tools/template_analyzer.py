"""
Template Analyzer Tool
=====================

AI-powered template analysis to extract fields and learn templates.
Uses PostgreSQL database for storage.
"""

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional, List, Dict

import httpx
from docx import Document
from psycopg import connect
from db.connection import get_db_conn


def get_db_connection():
    """Get database connection."""
    import os

    return get_db_conn()


def extract_placeholders(template_path: Path) -> dict[str, Any]:
    """
    Extract all placeholders from a Word template.
    Returns required fields and their locations.
    """
    doc = Document(str(template_path))

    placeholders = {}
    placeholder_pattern = re.compile(r"\{\{([^}]+)\}\}|\{([^}]+)\}|\[([^\]]+)\]")

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        matches = placeholder_pattern.findall(text)
        for match in matches:
            placeholder = (match[0] or match[1] or match[2]).strip()
            if placeholder and placeholder.lower() not in placeholders:
                placeholders[placeholder.lower()] = {
                    "field": placeholder,
                    "location": f"paragraph {idx + 1}",
                    "sample_text": text[:100] if text else "",
                }

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                matches = placeholder_pattern.findall(text)
                for match in matches:
                    placeholder = (match[0] or match[1] or match[2]).strip()
                    if placeholder and placeholder.lower() not in placeholders:
                        placeholders[placeholder.lower()] = {
                            "field": placeholder,
                            "location": f"table {table_idx + 1}, row {row_idx + 1}, cell {cell_idx + 1}",
                            "sample_text": text[:100] if text else "",
                        }

    return {
        "template": template_path.name,
        "total_placeholders": len(placeholders),
        "fields": list(placeholders.keys()),
        "details": placeholders,
    }


def classify_template_fields(template_text: str, fields: List[str]) -> Dict[str, Any]:
    """
    Use Claude Haiku via OpenRouter to classify each placeholder field as
    db_field, user_input, or flag static text warnings.

    Falls back to classifying all fields as user_input if the API call fails.
    """
    # Build the classification prompt
    fields_str = ", ".join(fields)
    truncated_text = template_text[:3000]

    prompt = (
        "Given this legal document template with these placeholder fields: "
        + fields_str
        + "\n\nTemplate text content:\n"
        + truncated_text
        + "\n\n"
        "Classify each placeholder field into one of these categories:\n"
        '1. "db_field" - Information that comes from a company database '
        "(company name, registration number, address, directors list, shareholders, etc.)\n"
        '2. "user_input" - Information that is specific to each transaction and must be '
        "entered by the user each time (specific director being appointed/resigning, "
        "date of action, NRC/passport number of specific person, phone, email, date of birth, etc.)\n\n"
        "Also identify any hardcoded/static text in the template that might need "
        "customization for different companies or jurisdictions.\n\n"
        "Return JSON:\n"
        "{\n"
        '  "db_fields": ["field1", "field2"],\n'
        '  "user_input_fields": ["field3", "field4"],\n'
        '  "field_descriptions": {"field1": "Company\'s registered name", "field3": "Name of director being appointed"},\n'
        '  "static_text_warnings": ["warning1", "warning2"]\n'
        "}\n\n"
        "Return ONLY JSON."
    )

    openrouter_key = os.getenv("OPENROUTER_API_KEY")
    if not openrouter_key:
        print("OPENROUTER_API_KEY not set, falling back to default classification")
        return _fallback_classification(fields)

    try:
        from app.model_config import OPENROUTER_BASE_URL
        response = httpx.post(
            f"{OPENROUTER_BASE_URL}/chat/completions",
            headers={
                "Authorization": f"Bearer {openrouter_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": os.getenv("CLASSIFICATION_MODEL", "google/gemini-3.1-flash-lite-preview"),
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0,
            },
            timeout=60,
        )
        response.raise_for_status()

        result = response.json()
        content = result.get("choices", [{}])[0].get("message", {}).get("content", "")

        # Parse the JSON from the response (strip markdown fences if present)
        content = content.strip()
        if content.startswith("```"):
            # Remove ```json ... ``` wrapper
            content = re.sub(r"^```(?:json)?\s*", "", content)
            content = re.sub(r"\s*```$", "", content)

        classification = json.loads(content)

        # Validate the structure and ensure all fields are accounted for
        db_fields = classification.get("db_fields", [])
        user_input_fields = classification.get("user_input_fields", [])
        field_descriptions = classification.get("field_descriptions", {})
        static_text_warnings = classification.get("static_text_warnings", [])

        # Make sure every field is classified somewhere
        classified = set(db_fields) | set(user_input_fields)
        for field in fields:
            if field not in classified:
                user_input_fields.append(field)

        return {
            "db_fields": db_fields,
            "user_input_fields": user_input_fields,
            "field_descriptions": field_descriptions,
            "static_text_warnings": static_text_warnings,
        }

    except Exception as e:
        print(f"Error classifying fields via OpenRouter: {e}")
        return _fallback_classification(fields)


def _fallback_classification(fields: List[str]) -> Dict[str, Any]:
    """Fallback classification when AI is unavailable — uses heuristic rules."""
    # Common DB fields (company-level data)
    DB_KEYWORDS = {
        "company", "company_name", "registration", "registration_number",
        "company_registration_number", "registered_office", "address",
        "directors", "shareholders", "total_shares", "total_capital",
        "currency", "date_of_incorporation", "company_type",
        "principal_place_of_business", "group_company", "foreign_company",
        "individual_shareholder", "corporate_shareholder",
    }

    db_fields = []
    user_input_fields = []

    for field in fields:
        field_lower = field.lower().replace(" ", "_")
        is_db = False
        for kw in DB_KEYWORDS:
            if kw in field_lower:
                is_db = True
                break
        if is_db:
            db_fields.append(field)
        else:
            user_input_fields.append(field)

    return {
        "db_fields": db_fields,
        "user_input_fields": user_input_fields,
        "field_descriptions": {},
        "static_text_warnings": [],
    }


def _get_template_text(template_path: Path) -> str:
    """Extract full text from a Word template for classification."""
    doc = Document(str(template_path))
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def save_template_to_db(name: str, path: str, fields: List[str], details: Dict, document_type: str = None) -> bool:
    """Save template to database with all fields."""
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Extract enhanced fields from details
        purpose = details.get("purpose", "")
        when_to_use = details.get("when_to_use", "")
        how_to_use = json.dumps(details.get("how_to_use", []))
        prerequisites = json.dumps(details.get("prerequisites", []))
        filing_deadline = details.get("filing_deadline", "")
        fees = details.get("fees", "")
        validity_period = details.get("validity_period", "")
        approval_chain = json.dumps(details.get("approval_chain", []))
        required_attachments = json.dumps(details.get("required_attachments", []))
        common_mistakes = json.dumps(details.get("common_mistakes", []))
        jurisdiction = details.get("jurisdiction", "Myanmar")
        industry_tags = json.dumps(details.get("industry_tags", []))
        complexity = details.get("complexity", "Medium")
        estimated_time = details.get("estimated_time", "")
        category = details.get("category", "")
        keywords = details.get("keywords", "")
        description = details.get("description", "")
        usage_instructions = details.get("usage_instructions", "")
        sections = json.dumps(details.get("sections", []))
        signatures = json.dumps(details.get("signatures", []))
        deadlines = json.dumps(details.get("deadlines", []))
        legal_references = json.dumps(details.get("legal_references", []))
        related_documents = json.dumps(details.get("related_documents", []))
        tips = json.dumps(details.get("tips", []))
        use_cases = json.dumps(details.get("use_cases", []))

        cur.execute(
            """
            INSERT INTO templates (
                name, path, fields, total_fields, document_type,
                purpose, when_to_use, how_to_use, prerequisites,
                filing_deadline, fees, validity_period, approval_chain,
                required_attachments, common_mistakes, jurisdiction,
                industry_tags, complexity, estimated_time,
                category, keywords, description, usage_instructions,
                sections, signatures, deadlines, legal_references,
                related_documents, tips, use_cases,
                ai_trained, updated_at
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (name) DO UPDATE SET
                path = EXCLUDED.path,
                fields = EXCLUDED.fields,
                total_fields = EXCLUDED.total_fields,
                document_type = EXCLUDED.document_type,
                purpose = EXCLUDED.purpose,
                when_to_use = EXCLUDED.when_to_use,
                how_to_use = EXCLUDED.how_to_use,
                prerequisites = EXCLUDED.prerequisites,
                filing_deadline = EXCLUDED.filing_deadline,
                fees = EXCLUDED.fees,
                validity_period = EXCLUDED.validity_period,
                approval_chain = EXCLUDED.approval_chain,
                required_attachments = EXCLUDED.required_attachments,
                common_mistakes = EXCLUDED.common_mistakes,
                jurisdiction = EXCLUDED.jurisdiction,
                industry_tags = EXCLUDED.industry_tags,
                complexity = EXCLUDED.complexity,
                estimated_time = EXCLUDED.estimated_time,
                category = EXCLUDED.category,
                keywords = EXCLUDED.keywords,
                description = EXCLUDED.description,
                usage_instructions = EXCLUDED.usage_instructions,
                sections = EXCLUDED.sections,
                signatures = EXCLUDED.signatures,
                deadlines = EXCLUDED.deadlines,
                legal_references = EXCLUDED.legal_references,
                related_documents = EXCLUDED.related_documents,
                tips = EXCLUDED.tips,
                use_cases = EXCLUDED.use_cases,
                ai_trained = EXCLUDED.ai_trained,
                updated_at = EXCLUDED.updated_at
        """,
            (
                name,
                path,
                json.dumps(details),
                len(fields),
                document_type,
                purpose,
                when_to_use,
                how_to_use,
                prerequisites,
                filing_deadline,
                fees,
                validity_period,
                approval_chain,
                required_attachments,
                common_mistakes,
                jurisdiction,
                industry_tags,
                complexity,
                estimated_time,
                category,
                keywords,
                description,
                usage_instructions,
                sections,
                signatures,
                deadlines,
                legal_references,
                related_documents,
                tips,
                use_cases,
                True,
                datetime.now(),
            ),
        )

        conn.commit()
        cur.close()
        return True
    except Exception as e:
        logging.getLogger("legalscout").warning(f"DB error in save_template_to_db: {e}")
        return False
    finally:
        if conn:
            conn.close()


def get_template_from_db(name: str) -> Optional[Dict]:
    """Get template from database."""
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT name, path, fields, total_fields, document_type, created_at, updated_at FROM templates WHERE name = %s",
            (name,),
        )
        row = cur.fetchone()
        cur.close()

        if row:
            return {
                "name": row[0],
                "path": row[1],
                "fields": row[2] if row[2] else {},
                "total_fields": row[3],
                "document_type": row[4],
                "created_at": row[5].isoformat() if row[5] else None,
                "updated_at": row[6].isoformat() if row[6] else None,
            }
        return None
    except Exception as e:
        logging.getLogger("legalscout").warning(f"DB error in get_template_from_db: {e}")
        return None
    finally:
        if conn:
            conn.close()


def get_all_templates_from_db() -> List[Dict]:
    """Get all templates from database."""
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT name, path, fields, total_fields, document_type, 
                   category, keywords, description,
                   purpose, when_to_use, how_to_use, prerequisites,
                   filing_deadline, fees, validity_period, approval_chain,
                   required_attachments, common_mistakes, jurisdiction,
                   industry_tags, complexity, estimated_time,
                   sections, signatures, deadlines, legal_references,
                   related_documents, tips, use_cases, usage_instructions,
                   ai_trained, created_at, updated_at, ai_analyzed, uploaded_by_email,
                   field_mapping
            FROM templates ORDER BY created_at DESC
        """)
        rows = cur.fetchall()
        cur.close()

        return [
            {
                "name": row[0],
                "path": row[1],
                "fields": row[2] if row[2] else {},
                "total_fields": row[3],
                "document_type": row[4],
                "category": row[5],
                "keywords": row[6],
                "description": row[7],
                "purpose": row[8],
                "when_to_use": row[9],
                "how_to_use": row[10],
                "prerequisites": row[11],
                "filing_deadline": row[12],
                "fees": row[13],
                "validity_period": row[14],
                "approval_chain": row[15],
                "required_attachments": row[16],
                "common_mistakes": row[17],
                "jurisdiction": row[18],
                "industry_tags": row[19],
                "complexity": row[20],
                "estimated_time": row[21],
                "sections": row[22],
                "signatures": row[23],
                "deadlines": row[24],
                "legal_references": row[25],
                "related_documents": row[26],
                "tips": row[27],
                "use_cases": row[28],
                "usage_instructions": row[29],
                "ai_trained": row[30],
                "created_at": row[31].isoformat() if row[31] else None,
                "updated_at": row[32].isoformat() if row[32] else None,
                "ai_analyzed": row[33] if len(row) > 33 else None,
                "uploaded_by_email": row[34] if len(row) > 34 else None,
                "field_mapping": row[35] if len(row) > 35 else None,
            }
            for row in rows
        ]
    except Exception as e:
        logging.getLogger("legalscout").warning(f"DB error in get_all_templates_from_db: {e}")
        return []
    finally:
        if conn:
            conn.close()


def analyze_template(template_name: str, documents_dir: str = "/documents") -> dict[str, Any]:
    """
    Analyze a template and extract all fields.
    Also saves to the database.
    """
    templates_dir = Path(documents_dir) / "legal" / "templates"
    template_path = templates_dir / template_name

    # If not found, try with spaces/underscores conversion
    if not template_path.exists():
        # Try converting underscores to spaces
        alt_name = template_name.replace("_", " ")
        template_path = templates_dir / alt_name
        if not template_path.exists():
            # Try converting spaces to underscores
            alt_name = template_name.replace(" ", "_")
            template_path = templates_dir / alt_name
            if not template_path.exists():
                return {
                    "success": False,
                    "error": f"Template not found: {template_name}",
                }
        # Use the found name
        template_name = alt_name if template_path.exists() else template_name

    result = extract_placeholders(template_path)

    document_type = _infer_document_type(template_name)

    # Classification happens during "Start Training", not during upload/analyze

    save_template_to_db(
        name=template_name,
        path=str(template_path),
        fields=result["fields"],
        details=result["details"],
        document_type=document_type,
    )

    return {
        "success": True,
        **result,
    }


def get_template_info(template_name: str) -> Optional[dict]:
    """Get saved template info from database."""
    return get_template_from_db(template_name)


def list_analyzed_templates() -> list:
    """List all analyzed templates."""
    templates = get_all_templates_from_db()
    return [
        {
            "name": t["name"],
            "path": t.get("path"),
            "fields": list(t["fields"].keys()) if isinstance(t["fields"], dict) else t.get("fields", []),
            "total_fields": t.get("total_fields", 0),
            "category": t.get("category"),
            "keywords": t.get("keywords"),
            "description": t.get("description"),
            "purpose": t.get("purpose"),
            "when_to_use": t.get("when_to_use"),
            "how_to_use": t.get("how_to_use"),
            "prerequisites": t.get("prerequisites"),
            "filing_deadline": t.get("filing_deadline"),
            "fees": t.get("fees"),
            "validity_period": t.get("validity_period"),
            "approval_chain": t.get("approval_chain"),
            "required_attachments": t.get("required_attachments"),
            "common_mistakes": t.get("common_mistakes"),
            "jurisdiction": t.get("jurisdiction"),
            "industry_tags": t.get("industry_tags"),
            "complexity": t.get("complexity"),
            "estimated_time": t.get("estimated_time"),
            "sections": t.get("sections"),
            "signatures": t.get("signatures"),
            "deadlines": t.get("deadlines"),
            "legal_references": t.get("legal_references"),
            "related_documents": t.get("related_documents"),
            "tips": t.get("tips"),
            "use_cases": t.get("use_cases"),
            "usage_instructions": t.get("usage_instructions"),
            "ai_trained": t.get("ai_trained"),
            "ai_analyzed": t.get("ai_analyzed"),
            "uploaded_by_email": t.get("uploaded_by_email"),
            "field_mapping": t.get("field_mapping"),
            "created_at": t.get("created_at"),
            "updated_at": t.get("updated_at"),
        }
        for t in templates
    ]


def save_template_knowledge(template_name: str, fields: list, documents_dir: str = "/documents") -> dict:
    """Save template knowledge to database."""
    template = get_template_from_db(template_name)

    if template:
        return {"success": True, "knowledge_saved": True}

    templates_dir = Path(documents_dir) / "legal" / "templates"
    template_path = templates_dir / template_name

    if template_path.exists():
        analyze_template(template_name, documents_dir)

    return {"success": True, "knowledge_saved": True}


def _infer_document_type(template_name: str) -> str:
    """Infer document type from template name."""
    name_lower = template_name.lower()

    if "agm" in name_lower or "annual" in name_lower:
        return "AGM Notice"
    elif "board" in name_lower:
        return "Board Meeting"
    elif "contract" in name_lower:
        return "Contract"
    elif "nda" in name_lower:
        return "NDA"
    elif "invoice" in name_lower:
        return "Invoice"
    elif "letter" in name_lower:
        return "Letter"
    else:
        return "Legal Document"


def create_template_analyzer_tool(documents_dir: str = "/documents"):
    """Create the template analyzer tool for the agent."""

    def analyze(template_name: str) -> dict[str, Any]:
        """Analyze a template and extract fields."""
        return analyze_template(template_name, documents_dir)

    def get_info(template_name: str) -> dict:
        """Get template information."""
        info = get_template_info(template_name)
        if info:
            return {"success": True, "template": info}
        return {"success": False, "error": "Template not found"}

    def list_templates() -> dict:
        """List all analyzed templates."""
        return {"templates": list_analyzed_templates()}

    def save_knowledge(template_name: str) -> dict:
        """Save template to agent knowledge base."""
        info = get_template_info(template_name)
        if not info:
            return {"success": False, "error": "Template not analyzed yet"}

        return save_template_knowledge(template_name, info.get("fields", []), documents_dir)

    return {
        "analyze_template": analyze,
        "get_template_info": get_info,
        "list_templates": list_templates,
        "save_template_knowledge": save_knowledge,
    }
