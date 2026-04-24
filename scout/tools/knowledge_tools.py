"""
Knowledge Base Tools for Agent
===========================

Tools for agent to search and use knowledge base data
for template auto-filling.
"""

import json
import logging
from typing import Any, Dict, List, Optional
from db.connection import get_db_conn
from scout.tools.knowledge_base import (
    search_knowledge,
    lookup_value,
    get_source_data,
    get_knowledge_sources,
)


def create_knowledge_tools():
    """Create knowledge base tools for the agent."""

    def search(query: str, limit: int = 10) -> Dict[str, Any]:
        """
        Search the knowledge base for information.

        Args:
            query: Search query (company name, person, etc.)
            limit: Number of results to return

        Returns:
            List of matching records with key-value pairs
        """
        results = search_knowledge(query, limit)
        return {"query": query, "results": results, "count": len(results)}

    def lookup(key: str, value: str) -> Dict[str, Any]:
        """
        Exact lookup - find record by key and value.

        Args:
            key: Field name to look up (e.g., "company_name", "name")
            value: Value to find (e.g., "CityHolding", "John Doe")

        Returns:
            Matching records with all fields
        """
        results = lookup_value(key, value)
        return {"key": key, "value": value, "results": results, "count": len(results)}

    def get_company(company_name: str) -> Dict[str, Any]:
        """Get company details — searches companies table directly with fuzzy match."""
        conn = None
        try:
            conn = get_db_conn()
            cur = conn.cursor()
            cur.execute("""
                SELECT company_name_english, company_registration_number, status, company_type,
                       principal_activity, registered_office_address, directors, members,
                       total_shares_issued, currency_of_share_capital, date_of_last_annual_return,
                       ultimate_holding_company_name, foreign_company
                FROM companies WHERE company_name_english ILIKE %s
                LIMIT 1
            """, (f"%{company_name}%",))
            row = cur.fetchone()
            cur.close()
            if row:
                dirs = row[6] if isinstance(row[6], list) else []
                mems = row[7] if isinstance(row[7], list) else []
                return {
                    "company": row[0], "found": True,
                    "data": {
                        "company_name": row[0], "registration_number": row[1],
                        "status": row[2], "company_type": row[3],
                        "principal_activity": row[4], "registered_office": row[5],
                        "directors": [d.get("name", str(d)) if isinstance(d, dict) else str(d) for d in dirs],
                        "shareholders": [m.get("name", str(m)) if isinstance(m, dict) else str(m) for m in mems],
                        "total_shares": row[8], "currency": row[9],
                        "last_annual_return": str(row[10]) if row[10] else None,
                        "holding_company": row[11], "foreign_company": row[12],
                    }
                }
        except Exception as e:
            logging.getLogger("legalscout").warning(f"DB error in get_company: {e}")
        finally:
            if conn:
                conn.close()
        # Fallback to knowledge_lookup
        results = lookup_value("company_name", company_name)
        if not results:
            results = lookup_value("name", company_name)
        return {"company": company_name, "found": len(results) > 0, "data": results[0] if results else None}

    def get_directors(company_name: str) -> List[Dict[str, Any]]:
        """Get directors for a company — searches companies table directly."""
        conn = None
        try:
            conn = get_db_conn()
            cur = conn.cursor()
            cur.execute("SELECT company_name_english, directors FROM companies WHERE company_name_english ILIKE %s LIMIT 1", (f"%{company_name}%",))
            row = cur.fetchone()
            cur.close()
            if row and row[1]:
                dirs = row[1] if isinstance(row[1], list) else []
                return [{"company": row[0], "data": d} for d in dirs]
        except Exception as e:
            logging.getLogger("legalscout").warning(f"DB error in get_directors: {e}")
        finally:
            if conn:
                conn.close()
        return []

    def get_shareholders(company_name: str) -> List[Dict[str, Any]]:
        """Get shareholders for a company — searches companies table directly."""
        conn = None
        try:
            conn = get_db_conn()
            cur = conn.cursor()
            cur.execute("SELECT company_name_english, members FROM companies WHERE company_name_english ILIKE %s LIMIT 1", (f"%{company_name}%",))
            row = cur.fetchone()
            cur.close()
            if row and row[1]:
                mems = row[1] if isinstance(row[1], list) else []
                return [{"company": row[0], "data": m} for m in mems]
        except Exception as e:
            logging.getLogger("legalscout").warning(f"DB error in get_shareholders: {e}")
        finally:
            if conn:
                conn.close()
        return []

    def get_template_data(template_name: str) -> Dict[str, Any]:
        """
        Get all data relevant for a template.
        Analyzes template placeholders and finds matching data.

        Args:
            template_name: Name of the template

        Returns:
            All relevant data for filling this template
        """
        # Get template info to see what fields are needed
        from scout.tools.template_analyzer import get_template_from_db

        template_info = get_template_from_db(template_name)
        required_fields = template_info.get("fields", {}) if template_info else {}

        data = {
            "template": template_name,
            "required_fields": list(required_fields.keys()) if required_fields else [],
            "matched_data": {},
        }

        # Try to find matching data for each field
        for field_name in required_fields.keys():
            # Search for this field
            results = search_knowledge(field_name, 5)
            if results:
                data["matched_data"][field_name] = results

        return data

    def list_sources() -> List[str]:
        """List all available knowledge sources."""
        sources = get_knowledge_sources()
        return [s["filename"] for s in sources]

    def get_data_for_template(company_name: str) -> Dict[str, Any]:
        """
        Get all data needed to fill a template for a company.

        Args:
            company_name: Name of the company

        Returns:
            Dictionary with all data to fill template placeholders
        """
        # Get company info
        company_info = lookup_value("company_name", company_name)
        if not company_info:
            company_info = lookup_value("name", company_name)

        # Get directors
        directors = get_directors(company_name)

        # Get shareholders
        shareholders = get_shareholders(company_name)

        return {
            "company_name": company_name,
            "company_info": company_info[0]["data"] if company_info else None,
            "directors": [d["data"] for d in directors],
            "shareholders": [s["data"] for s in shareholders],
        }

    def generate_dica_extract(company_name: str) -> Dict[str, Any]:
        """Generate a DICA Company Extract document for a company.
        Creates a .docx file with all company details formatted as an official DICA extract.
        Returns download link and company summary."""
        import os, json
        conn = None
        try:
            from psycopg import connect
            conn = get_db_conn()
            cur = conn.cursor()
            cur.execute("""
                SELECT id, company_name_english, company_name_myanmar, company_registration_number,
                       registration_date, status, company_type, foreign_company,
                       principal_activity, registered_office_address, principal_place_of_business,
                       directors, total_shares_issued, currency_of_share_capital, members,
                       date_of_last_annual_return, ultimate_holding_company_name, filing_history
                FROM companies WHERE company_name_english ILIKE %s
            """, (f"%{company_name}%",))
            row = cur.fetchone()
            cur.close()

            if not row:
                return {"success": False, "error": f"Company '{company_name}' not found in database"}

            company_id = row[0]
            name = row[1] or company_name
            dirs = row[11] if isinstance(row[11], list) else []
            members = row[14] if isinstance(row[14], list) else []

            # Generate DICA extract document
            from docx import Document
            from pathlib import Path
            import datetime

            doc = Document()
            doc.add_heading("DICA Company Extract", 0)
            doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%d %B %Y, %H:%M')}")
            doc.add_paragraph("")

            # Company details
            doc.add_heading("Company Information", level=1)
            details = [
                ("Company Name (English)", row[1] or ""),
                ("Company Name (Myanmar)", row[2] or ""),
                ("Registration Number", row[3] or ""),
                ("Registration Date", str(row[4]) if row[4] else ""),
                ("Status", row[5] or ""),
                ("Company Type", row[6] or ""),
                ("Foreign Company", row[7] or ""),
                ("Principal Activity", row[8] or ""),
                ("Registered Office", row[9] or ""),
                ("Principal Place of Business", row[10] or ""),
                ("Total Shares Issued", row[12] or ""),
                ("Currency", row[13] or ""),
                ("Last Annual Return", str(row[15]) if row[15] else ""),
                ("Ultimate Holding Company", row[16] or ""),
            ]
            table = doc.add_table(rows=len(details), cols=2)
            table.style = "Table Grid"
            for i, (label, value) in enumerate(details):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = str(value)

            # Directors
            if dirs:
                doc.add_heading("Directors", level=1)
                for d in dirs:
                    if isinstance(d, dict):
                        doc.add_paragraph(f"• {d.get('name', '')} — {d.get('position', 'Director')}")
                    else:
                        doc.add_paragraph(f"• {d}")

            # Members/Shareholders
            if members:
                doc.add_heading("Members / Shareholders", level=1)
                for m in members:
                    if isinstance(m, dict):
                        shares = m.get('shares', '')
                        doc.add_paragraph(f"• {m.get('name', '')} — {shares} shares")
                    else:
                        doc.add_paragraph(f"• {m}")

            # Save
            safe_name = name.replace(" ", "_").replace("/", "_")[:50]
            filename = f"DICA_Extract_{safe_name}.docx"
            output_dir = Path("/documents/legal/output")
            output_dir.mkdir(parents=True, exist_ok=True)
            filepath = output_dir / filename
            doc.save(str(filepath))

            download_url = f"/documents/legal/output/{filename}"

            # Build summary
            director_names = [d.get("name", str(d)) if isinstance(d, dict) else str(d) for d in dirs[:5]]
            member_names = [m.get("name", str(m)) if isinstance(m, dict) else str(m) for m in members[:5]]

            return {
                "success": True,
                "company_name": name,
                "registration_number": row[3] or "",
                "status": row[5] or "",
                "company_type": row[6] or "",
                "directors": director_names,
                "shareholders": member_names,
                "registered_office": row[9] or "",
                "total_shares": row[12] or "",
                "file_name": filename,
                "download_url": download_url,
                "message": f"DICA Extract generated for {name}",
            }
        except Exception as e:
            logging.getLogger("legalscout").warning(f"DB error in generate_dica_extract: {e}")
            return {"success": False, "error": str(e)}
        finally:
            if conn:
                conn.close()

    return {
        "search_knowledge": search,
        "lookup_knowledge": lookup,
        "get_company": get_company,
        "get_directors": get_directors,
        "get_shareholders": get_shareholders,
        "get_template_data": get_template_data,
        "list_knowledge_sources": list_sources,
        "get_data_for_template": get_data_for_template,
        "generate_dica_extract": generate_dica_extract,
    }
