"""
Smart Document Workflow Tool
===========================

Complete workflow for legal document generation:
1. Analyze template - extract required fields
2. Check data - validate company data vs required fields
3. Ask for missing info - if needed
4. Validate filled document
5. Generate final document
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

# Context for current document generation (set by generate_document, used by find_replacement)
_current_template_name: str | None = None
_current_company_name: str | None = None

from scout.tools.document_tracker import record_document


def extract_placeholders_from_template(template_path: Path) -> dict[str, Any]:
    """
    Extract all placeholders from a Word template.
    Returns required fields and their locations.
    """
    doc = Document(str(template_path))

    placeholders = {}
    placeholder_pattern = re.compile(r"\{\{([^}]+)\}\}|\{([^}]+)\}|\[([^\]]+)\]")

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        matches = placeholder_pattern.findall(text)
        for match in matches:
            placeholder = (match[0] or match[1] or match[2]).strip()
            if placeholder:
                placeholders[placeholder.lower()] = {
                    "field": placeholder,
                    "location": f"paragraph {idx + 1}",
                    "sample_text": text[:100] if text else "",
                }

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                matches = placeholder_pattern.findall(text)
                for match in matches:
                    placeholder = (match[0] or match[1] or match[2]).strip()
                    if placeholder:
                        placeholders[placeholder.lower()] = {
                            "field": placeholder,
                            "location": f"table {table_idx + 1}, row {row_idx + 1}, cell {cell_idx + 1}",
                            "sample_text": text[:100] if text else "",
                        }

    return {
        "template": template_path.name,
        "total_placeholders": len(placeholders),
        "fields": list(placeholders.keys()),
        "details": placeholders,
    }



def validate_data_vs_template(required_fields: list, company_data: dict) -> dict[str, Any]:
    """
    Validate that company data has all required fields for template.
    Returns what's available vs what's missing.
    """
    available_columns = [col.lower().replace(" ", "_").replace("-", "_") for col in company_data.get("columns", list(company_data.keys()))]

    available_fields = []
    missing_fields = []

    # Fields that always have defaults
    DEFAULT_FIELDS = {"pronoun", "meeting_location", "financial_year_end_date",
                      "next_financial_year_end_date", "auditor_name", "auditor_fee",
                      "date", "address", "nric", "identification_type",
                      "authorized_director_name", "resigning_director_name",
                      "resigning_director_identification_number", "shareholder_name"}

    # Alias mapping for matching
    FIELD_ALIASES = {
        "company": "company_name", "date": "meeting_date",
        "meeting_location": "registered_office",
        "director_name": "directors", "authorized_director_name": "directors",
        "resigning_director_name": "directors", "address": "registered_office",
        "nric": "nrc_passport", "identification_type": "nrc_passport",
        "shareholder_name": "individual_shareholder_1_name",
    }

    for field in required_fields:
        field_normalized = field.lower().replace(" ", "_").replace("-", "_")

        found = False
        # Direct match
        for col in available_columns:
            if field_normalized in col or col in field_normalized:
                found = True
                break

        # Alias match
        if not found and field_normalized in FIELD_ALIASES:
            alias = FIELD_ALIASES[field_normalized]
            for col in available_columns:
                if alias in col or col in alias:
                    found = True
                    break

        # Has default
        if not found and field_normalized in DEFAULT_FIELDS:
            found = True

        if found:
            available_fields.append(field)
        else:
            missing_fields.append(field)

    return {
        "required_fields": required_fields,
        "available_fields": available_fields,
        "missing_fields": missing_fields,
        "available_columns": available_columns,
        "is_complete": len(missing_fields) == 0,
    }


def find_company_data(company_name: str, companies_data: dict) -> dict[str, Any]:
    """Find specific company data from companies dict."""
    company_name_lower = company_name.lower().strip()

    for company in companies_data.get("companies", []):
        comp_name = str(company.get("company_name", "")).lower().strip()
        if comp_name == company_name_lower or company_name_lower in comp_name:
            return {"found": True, "data": company}

    return {"found": False, "available_companies": [c.get("company_name") for c in companies_data.get("companies", [])]}


def analyze_template(template_name: str, documents_dir: str = "/documents") -> dict[str, Any]:
    """
    Step 1: Analyze template to extract required fields.
    """
    templates_dir = Path(documents_dir) / "legal" / "templates"
    template_path = templates_dir / template_name

    if not template_path.exists():
        return {
            "success": False,
            "error": f"Template not found: {template_name}",
            "available_templates": [f.name for f in templates_dir.glob("*.docx")] if templates_dir.exists() else [],
        }

    result = extract_placeholders_from_template(template_path)
    result["success"] = True
    result["template_path"] = str(template_path)

    return result


def prepare_document_data(template_name: str, company_name: str, documents_dir: str = "/documents") -> dict[str, Any]:
    """
    Complete workflow: Analyze → Validate → Prepare data for document.

    Returns:
        - Template analysis (required fields)
        - Company data from database
        - Validation (what's missing)
        - Ready status
    """
    templates_dir = Path(documents_dir) / "legal" / "templates"

    template_path = templates_dir / template_name
    if not template_path.exists():
        # Try converting underscores to spaces
        alt_name = template_name.replace("_", " ")
        template_path = templates_dir / alt_name
        if not template_path.exists():
            # Try converting spaces to underscores
            alt_name = template_name.replace(" ", "_")
            template_path = templates_dir / alt_name
            if not template_path.exists():
                return {"success": False, "error": f"Template not found: {template_name}", "step": "template_analysis"}
        # Use the found name
        template_name = alt_name if template_path.exists() else template_name

    template_analysis = extract_placeholders_from_template(template_path)
    required_fields = template_analysis.get("fields", [])

    # Load company data from database
    company_result = None
    companies_data = {"companies": []}
    try:
        from scout.tools.companies_db import get_all_companies
        db_companies = get_all_companies(limit=200)
        if db_companies:
            companies_list = []
            for c in db_companies:
                director_names = c.get("directors", "")
                shareholder_names = c.get("shareholders", "")
                shareholders_list = c.get("shareholders_list", [])
                dirs_split = [d.strip() for d in director_names.split(",") if d.strip()] if director_names else []
                shs_split = [s.strip() for s in shareholder_names.split(",") if s.strip()] if shareholder_names else []

                sh_shares = {}
                for m in shareholders_list:
                    if isinstance(m, dict):
                        sh_shares[m.get("name", "")] = m.get("shares", "TBD")

                companies_list.append({
                    "company_name": c.get("name", ""),
                    "company": c.get("name", ""),
                    "company_name_english": c.get("name", ""),
                    "company_registration_number": c.get("registration_number", ""),
                    "registered_office": c.get("address", ""),
                    "registered_office_address": c.get("address", ""),
                    "company_address": c.get("address", ""),
                    "address": c.get("address", ""),
                    "meeting_location": c.get("address", ""),
                    "directors": director_names,
                    "director_name": dirs_split[0] if dirs_split else "TBD",
                    "authorized_director_name": dirs_split[0] if dirs_split else "TBD",
                    "authorized director_name": dirs_split[0] if dirs_split else "TBD",
                    "individual_shareholder_1_name": shs_split[0] if len(shs_split) > 0 else "TBD",
                    "individual_shareholder_2_name": shs_split[1] if len(shs_split) > 1 else "TBD",
                    "individual shareholder_1_name": shs_split[0] if len(shs_split) > 0 else "TBD",
                    "individual shareholder_2_name": shs_split[1] if len(shs_split) > 1 else "TBD",
                    "corporate_shareholder_3_name": shs_split[2] if len(shs_split) > 2 else "TBD",
                    "corporate shareholder_3_name": shs_split[2] if len(shs_split) > 2 else "TBD",
                    "shareholder_1_name": shs_split[0] if len(shs_split) > 0 else "TBD",
                    "shareholder_2_name": shs_split[1] if len(shs_split) > 1 else "TBD",
                    "shareholder_1_shares": sh_shares.get(shs_split[0], "TBD") if shs_split else "TBD",
                    "shareholder_2_shares": sh_shares.get(shs_split[1], "TBD") if len(shs_split) > 1 else "TBD",
                    "status": c.get("status", ""),
                    "company_type": c.get("company_type", ""),
                    "total_shares": c.get("total_shares", ""),
                    "currency": c.get("currency", ""),
                })
            companies_data = {"success": True, "companies": companies_list}
            company_result = find_company_data(company_name, companies_data)
    except Exception:
        pass

    if not company_result:
        return {"success": False, "error": "No company data in database. Add companies from the admin panel.", "step": "data_loading", "required_fields": required_fields}

    if not company_result.get("found"):
        return {
            "success": False,
            "error": f"Company '{company_name}' not found in database",
            "step": "company_lookup",
            "available_companies": company_result.get("available_companies", []),
            "required_fields": required_fields,
        }

    company_data = company_result.get("data", {})
    validation = validate_data_vs_template(required_fields, company_data)

    # Get template info from database (trained data)
    template_info = {}
    try:
        from scout.tools.template_analyzer import get_template_info

        template_data = get_template_info(template_name)
        if template_data:
            template_info = {
                "purpose": template_data.get("purpose", ""),
                "when_to_use": template_data.get("when_to_use", ""),
                "how_to_use": template_data.get("how_to_use", []),
            }
    except Exception:
        pass

    normalized_company_data = {}
    for key, value in company_data.items():
        key_normalized = str(key).lower().replace(" ", "_").replace("-", "_")
        normalized_company_data[key_normalized] = value

    return {
        "success": True,
        "step": "complete",
        "template_analysis": {
            "template": template_name,
            "required_fields": required_fields,
            "total_fields": len(required_fields),
            **template_info,
        },
        "company_data": company_data,
        "normalized_data": normalized_company_data,
        "validation": validation,
        "ready_to_generate": validation["is_complete"],
        "message": "Ready to generate document"
        if validation["is_complete"]
        else f"Missing fields: {', '.join(validation['missing_fields'])}",
    }


def validate_filled_document(document_path: Path, all_placeholders: list = None) -> dict[str, Any]:
    """
    Step 4: Validate that a filled document has no unfilled placeholders.

    Args:
        document_path: Path to the filled document
        all_placeholders: List of all placeholders from the original template
    """
    if not document_path.exists():
        return {"success": False, "error": "Document not found"}

    doc = Document(str(document_path))
    placeholder_pattern = re.compile(r"\{\{([^}]+)\}\}|\{([^}]+)\}|\[([^\]]+)\]")

    # Find unfilled placeholders in the filled document
    unfilled_placeholders = []
    unfilled_locations = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        matches = placeholder_pattern.findall(text)
        for match in matches:
            placeholder = (match[0] or match[1] or match[2]).strip()
            if placeholder and placeholder.lower() not in ["", " ", "null", "none"]:
                unfilled_placeholders.append(placeholder)
                unfilled_locations.append(f"Paragraph: {text[:50]}...")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                matches = placeholder_pattern.findall(text)
                for match in matches:
                    placeholder = (match[0] or match[1] or match[2]).strip()
                    if placeholder and placeholder.lower() not in ["", " ", "null", "none"]:
                        unfilled_placeholders.append(placeholder)
                        unfilled_locations.append(f"Table: {text[:50]}...")

    # Calculate filled placeholders
    if all_placeholders:
        unfilled_set = set(p.lower() for p in unfilled_placeholders)
        filled_placeholders = [p for p in all_placeholders if p.lower() not in unfilled_set]
        total_placeholders = len(all_placeholders)
        total_filled = len(filled_placeholders)
    else:
        # Fallback if all_placeholders not provided
        filled_placeholders = []
        total_placeholders = len(unfilled_placeholders)
        total_filled = 0

    unique_unfilled = list(set(unfilled_placeholders))
    return {
        "success": True,
        "is_valid": len(unique_unfilled) == 0,
        "all_placeholders": all_placeholders or [],
        "filled_placeholders": filled_placeholders,
        "unfilled_placeholders": unfilled_locations,
        "unfilled_names": unique_unfilled,
        "total_placeholders": total_placeholders,
        "total_filled": total_filled,
        "total_unfilled": len(unique_unfilled),
        "message": f"Filled {total_filled}/{total_placeholders} fields"
        if all_placeholders
        else ("Document is valid - all fields filled" if len(unfilled_placeholders) == 0 else f"Warning: {len(unfilled_placeholders)} unfilled placeholders found"),
    }


def create_smart_document_tool(documents_dir: str = "/documents", host: str = ""):
    """Create the smart document workflow tool."""

    def analyze_template_tool(template_name: str) -> dict[str, Any]:
        """Analyze a template to see what fields are required."""
        return analyze_template(template_name, documents_dir)

    def prepare_document(template_name: str, company_name: str) -> dict[str, Any]:
        """Complete workflow: analyze, validate, prepare data."""
        return prepare_document_data(template_name, company_name, documents_dir)

    def generate_document(template_name: str, company_name: str, custom_data: dict = None) -> dict[str, Any]:
        """Generate document with validation workflow."""
        global _current_template_name, _current_company_name
        _current_template_name = template_name
        _current_company_name = company_name

        result = prepare_document_data(template_name, company_name, documents_dir)

        if not result.get("success"):
            return result

        # Smart default: If 90%+ fields available, proceed automatically
        validation = result.get("validation", {})
        # Use available_fields (correct key) instead of matched_fields
        total_fields = len(validation.get("available_fields", [])) + len(validation.get("missing_fields", []))
        matched_fields = len(validation.get("available_fields", []))

        # Calculate percentage
        if total_fields > 0:
            field_coverage = matched_fields / total_fields
        else:
            field_coverage = 0

        # Always auto-proceed — use defaults/TBD for missing fields
        auto_proceed = True

        # --- Load field classification from DB ---
        field_classification = None
        try:
            from scout.tools.template_analyzer import get_template_from_db
            tpl_data = get_template_from_db(template_name)
            if tpl_data and isinstance(tpl_data.get("fields"), dict):
                field_classification = tpl_data["fields"].get("field_classification")
        except Exception:
            pass

        # If classification exists, separate user_input vs db fields
        if field_classification and not custom_data:
            user_input_fields = field_classification.get("user_input_fields", [])
            db_fields_list = field_classification.get("db_fields", [])
            field_descriptions = field_classification.get("field_descriptions", {})

            # Check which user_input_fields are actually missing from data
            normalized_data = result.get("normalized_data", {})
            missing_user_fields = []
            for uf in user_input_fields:
                uf_norm = uf.lower().replace(" ", "_").replace("-", "_")
                if not normalized_data.get(uf_norm):
                    missing_user_fields.append(uf)

            # Auto-filled DB fields
            db_fields_filled = []
            for df in db_fields_list:
                df_norm = df.lower().replace(" ", "_").replace("-", "_")
                if normalized_data.get(df_norm):
                    db_fields_filled.append(df)

            # Generate smart defaults for missing fields
            from datetime import datetime
            try:
                import zoneinfo, os as _dos
                _tz = zoneinfo.ZoneInfo(_dos.getenv("TZ", "Asia/Yangon"))
                _today = datetime.now(_tz).strftime("%Y-%m-%d")
            except Exception:
                _today = datetime.now().strftime("%Y-%m-%d")

            field_defaults = {}
            for mf in missing_user_fields:
                mf_lower = mf.lower()
                if "date" in mf_lower and "financial" not in mf_lower:
                    field_defaults[mf] = _today
                elif "location" in mf_lower or "meeting_location" in mf_lower:
                    field_defaults[mf] = normalized_data.get("registered_office_address", normalized_data.get("registered_office", "TBD"))
                elif "pronoun" in mf_lower:
                    field_defaults[mf] = "they"
                else:
                    field_defaults[mf] = "TBD"

            # If there are missing user input fields, return them for the chat to show a form
            if missing_user_fields:
                return {
                    "success": False,
                    "error": "Need user input for some fields",
                    "user_input_fields": missing_user_fields,
                    "field_descriptions": {
                        f: field_descriptions.get(f, "") for f in missing_user_fields
                    },
                    "field_defaults": field_defaults,
                    "db_fields_filled": db_fields_filled,
                    "static_text_warnings": field_classification.get("static_text_warnings", []),
                    "message": "These fields need your input (company data is already filled)",
                }

        # Auto-generate defaults for missing fields
        if not result.get("ready_to_generate") and not custom_data:
            custom_data = {}
            from datetime import datetime
            try:
                import zoneinfo, os
                tz = zoneinfo.ZoneInfo(os.getenv("TZ", "Asia/Yangon"))
                today = datetime.now(tz).strftime("%Y-%m-%d")
            except Exception:
                today = datetime.now().strftime("%Y-%m-%d")

            for field in validation.get("missing_fields", []):
                field_lower = field.lower()
                # Smart defaults for date fields
                if "date" in field_lower and "financial" not in field_lower:
                    custom_data[field] = today
                elif "meeting_location" in field_lower or "location" in field_lower:
                    custom_data[field] = result.get("normalized_data", {}).get("registered_office_address", "TBD")
                elif "pronoun" in field_lower:
                    custom_data[field] = "they"
                else:
                    custom_data[field] = "TBD"
            result["message"] = f"Auto-filling {len(validation.get('missing_fields', []))} fields with defaults"

        template_path = Path(documents_dir) / "legal" / "templates" / template_name
        if not template_path.exists():
            return {"success": False, "error": "Template not found"}

        data = result.get("normalized_data", {})
        if custom_data:
            data.update(custom_data)

        filled_doc = fill_template_with_validation(template_path, data)

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        # Clean company name - remove registration number, parentheses, etc.
        import re

        company_clean = re.sub(r"\([^)]*\)", "", company_name)  # Remove (Registration No: ...)
        company_clean = re.sub(r"\s+", " ", company_clean).strip()  # Clean extra spaces
        company_safe = (
            company_clean.replace(" ", "_").replace("/", "_").replace("'", "").replace("(", "").replace(")", "")
        )

        # Sanitize template name (remove spaces and special chars for URL-safe filename)
        template_safe = (
            template_name.replace('.docx', '')
            .replace(" ", "_")  # Replace spaces with underscores
            .replace("/", "_")
            .replace("\\", "_")
            .replace("(", "")
            .replace(")", "")
            .replace("[", "")
            .replace("]", "")
            .replace("'", "")
            .replace('"', "")
            .replace(",", "")
            .replace(":", "")
            .replace(";", "")
        )

        output_filename = f"{template_safe}_{company_safe}_{timestamp}.docx"
        output_dir = Path(documents_dir) / "legal" / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / output_filename

        filled_doc.save(str(output_path))

        # Upload to S3 (background)
        try:
            from app.s3_storage import s3_upload_async
            s3_upload_async(str(output_path))
        except Exception:
            pass

        # Extract placeholders from template BEFORE validation
        template_analysis = extract_placeholders_from_template(template_path)
        all_template_placeholders = template_analysis.get("fields", [])

        validation = validate_filled_document(output_path, all_template_placeholders)

        preview = generate_preview(filled_doc)

        doc_url = f"{host}/documents/legal/output/{output_filename}"
        google_viewer_url = f"https://docs.google.com/gview?embedded=1&url={doc_url}"

        record_document(
            template_name=template_name,
            company_name=company_name,
            file_name=output_filename,
            file_path=str(output_path),
            download_url=doc_url,
            preview_url=google_viewer_url,
            validation_result=validation,
            custom_data=data,
        )

        return {
            "success": True,
            "step": "generated",
            "file_name": output_filename,
            "file_path": str(output_path),
            "download_url": doc_url,
            "validation": validation,
            "preview": preview,
            "message": f"Document created successfully for {company_name}",
            "warnings": validation.get("unfilled_placeholders", []) if not validation.get("is_valid") else [],
            # Validation summary for user
            "validation_summary": {
                "total_placeholders": validation.get("total_placeholders", 0),
                "filled_from_data": validation.get("total_filled", 0),
                "unfilled": validation.get("total_unfilled", 0),
                "unfilled_fields": validation.get("unfilled_names", []),
                "validation_status": "Complete" if validation.get("is_valid") else "Partial - some fields may be empty",
            },
            # Template info from training
            "template_info": {
                "purpose": result.get("template_analysis", {}).get("purpose", ""),
                "when_to_use": result.get("template_analysis", {}).get("when_to_use", ""),
                "how_to_use": result.get("template_analysis", {}).get("how_to_use", []),
            },
        }

    # Combined tool: Analyze + Prepare + Generate in one call
    def create_document(template_name: str, company_name: str, custom_data: dict = None) -> dict:
        """
        Complete document creation in one call.
        Combines analyze, prepare, and generate.

        Args:
            template_name: Name of the template file
            company_name: Name of the company
            custom_data: Optional custom data to override database data

        Returns:
            Full result with download URL and validation summary
        """
        return generate_document(template_name, company_name, custom_data)

    def preview_doc(template_name: str, company_name: str, custom_data: dict = None) -> dict[str, Any]:
        """Preview document - shows what will be filled without saving."""

        result = prepare_document_data(template_name, company_name, documents_dir)

        if not result.get("success"):
            return result

        validation = result.get("validation", {})
        # Use available_fields (correct key) instead of matched_fields
        total_fields = len(validation.get("available_fields", [])) + len(validation.get("missing_fields", []))
        matched_fields = len(validation.get("available_fields", []))

        if total_fields > 0:
            field_coverage = matched_fields / total_fields
        else:
            field_coverage = 0

        # Build preview data
        preview_data = result.get("normalized_data", {}).copy()

        # Add custom data
        if custom_data:
            preview_data.update(custom_data)

        # Add TBD for missing fields
        for field in validation.get("missing_fields", []):
            if field not in preview_data:
                preview_data[field] = "[TBD - needs input]"

        # Get template info from database
        from scout.tools.template_analyzer import get_all_templates_from_db

        templates_db = get_all_templates_from_db()
        template_info = None
        for t in templates_db:
            if t.get("name") == template_name:
                template_info = t
                break

        # Generate table format preview
        preview_lines = []

        # Header
        preview_lines.append("📄 **DOCUMENT PREVIEW**")
        preview_lines.append("")

        # Template & Company info
        preview_lines.append(f"**Template:** {template_name}")
        preview_lines.append(f"**Company:** {company_name}")
        preview_lines.append(f"**Data Coverage:** {int(field_coverage * 100)}%")
        preview_lines.append("")

        # Template purpose (if available)
        if template_info and template_info.get("purpose"):
            preview_lines.append(f"📋 **Purpose:** {template_info['purpose']}")
            preview_lines.append("")

        # Data table
        preview_lines.append("### 📊 Data to be filled")
        preview_lines.append("")
        preview_lines.append("| Field | Value | Source |")
        preview_lines.append("|-------|-------|--------|")

        # Show matched fields (available_fields is the correct key)
        for field in validation.get("available_fields", []):
            value = preview_data.get(field, "[empty]")
            # Truncate long values
            if len(str(value)) > 40:
                value = str(value)[:37] + "..."
            preview_lines.append(f"| {field} | {value} | ✅ DB |")

        # Show missing fields
        for field in validation.get("missing_fields", []):
            value = preview_data.get(field, "[TBD]")
            preview_lines.append(f"| {field} | {value} | ⚠️ Needs |")

        preview_lines.append("")
        preview_lines.append("---")
        preview_lines.append("")
        preview_lines.append("**Ready to proceed?**")
        preview_lines.append("")
        # Extract short template name for button
        short_name = template_name.replace(".docx", "").replace("_", " ")
        if len(short_name) > 40:
            short_name = short_name[:37] + "..."
        preview_lines.append(f"a) ✅ Generate \"{short_name}\" for {company_name}")
        preview_lines.append("b) ❌ Cancel and modify data")
        preview_lines.append("")
        preview_lines.append("What would you like to do?")

        return {
            "success": True,
            "preview": "\n".join(preview_lines),
            "template_name": template_name,
            "company_name": company_name,
            "field_coverage": f"{int(field_coverage * 100)}%",
            "matched_fields": validation.get("available_fields", []),  # Use correct key
            "missing_fields": validation.get("missing_fields", []),
            "data": preview_data,
            "needs_approval": True,
        }

    return {
        "analyze_template": analyze_template_tool,
        "prepare_document": prepare_document,
        "generate_document": generate_document,
        "create_document": create_document,
        "preview_document": preview_doc,
    }


def _fill_paragraph_highlighted(paragraph, data: dict[str, Any], placeholder_pattern):
    """Fill placeholders in a paragraph, highlighting replaced text in yellow."""
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from copy import deepcopy
    import lxml.etree as etree

    text = paragraph.text
    matches = list(placeholder_pattern.finditer(text))
    if not matches:
        return

    # Get the formatting from the first run (to preserve font, size, etc.)
    base_run = paragraph.runs[0] if paragraph.runs else None

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # Build segments: [("text", highlight_color), ...]
    # highlight_color: None=no highlight, "yellow"=filled, "red"=TBD/unfilled
    segments = []
    last_end = 0
    for match in matches:
        placeholder = (match.group(1) or match.group(2) or match.group(3)).strip().lower()
        replacement = find_replacement(placeholder, data,
            template_name=_current_template_name, company_name=_current_company_name)

        if match.start() > last_end:
            segments.append((text[last_end:match.start()], None))

        if replacement:
            is_tbd = str(replacement).strip().upper() == "TBD"
            segments.append((str(replacement), "red" if is_tbd else "yellow"))
        else:
            segments.append((text[match.start():match.end()], None))

        last_end = match.end()

    if last_end < len(text):
        segments.append((text[last_end:], False))

    # Add runs for each segment
    p_element = paragraph._element
    # Remove all existing run elements
    for r in list(p_element.findall(qn('w:r'))):
        p_element.remove(r)

    for seg_text, highlight_color in segments:
        new_run = paragraph.add_run(seg_text)
        # Copy formatting from base run if available
        if base_run and base_run._element.find(qn('w:rPr')) is not None:
            new_rpr = deepcopy(base_run._element.find(qn('w:rPr')))
            # Remove any existing highlight from copied format
            for existing_hl in new_rpr.findall(qn('w:highlight')):
                new_rpr.remove(existing_hl)
            new_run._element.insert(0, new_rpr)

        # Add highlight: yellow=filled from data, red=TBD/unfilled
        if highlight_color:
            rpr = new_run._element.get_or_add_rPr()
            highlight = etree.SubElement(rpr, qn('w:highlight'))
            highlight.set(qn('w:val'), highlight_color)  # "yellow" for filled, "red" for TBD


def fill_template_with_validation(template_path: Path, data: dict[str, Any]) -> Document:
    """Fill template with data, highlighting filled values in yellow."""
    doc = Document(str(template_path))
    placeholder_pattern = re.compile(r"\{\{([^}]+)\}\}|\{([^}]+)\}|\[([^\]]+)\]")

    for paragraph in doc.paragraphs:
        _fill_paragraph_highlighted(paragraph, data, placeholder_pattern)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fill_paragraph_highlighted(paragraph, data, placeholder_pattern)

    return doc


def _get_company_field(company_row: dict, column_path: str) -> str | None:
    """Resolve a DB column path from company data.
    Handles: 'company_name_english', 'members[0].name', 'directors[1].position'"""
    if not company_row or not column_path:
        return None

    # Handle array access: members[0].name
    import re
    array_match = re.match(r'(\w+)\[(\d+)\]\.(\w+)', column_path)
    if array_match:
        field, idx, subfield = array_match.groups()
        arr = company_row.get(field, [])
        if isinstance(arr, list) and int(idx) < len(arr):
            item = arr[int(idx)]
            if isinstance(item, dict):
                return str(item.get(subfield, ""))
        return None

    # Direct field access
    return str(company_row.get(column_path, "")) if company_row.get(column_path) else None


def _get_field_mapping(template_name: str) -> dict | None:
    """Load learned field_mapping from DB for a template."""
    try:
        from db.connection import get_db_conn
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT field_mapping FROM templates WHERE name = %s", (template_name,))
        row = cur.fetchone()
        cur.close(); conn.close()
        if row and row[0]:
            return row[0] if isinstance(row[0], dict) else __import__('json').loads(row[0])
    except Exception:
        pass
    return None


def _get_company_from_db(company_name: str) -> dict | None:
    """Get full company record directly from companies table."""
    try:
        from db.connection import get_db_conn
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT company_name_english, company_registration_number, registered_office_address,
                   principal_place_of_business, status, company_type, directors, members,
                   total_shares_issued, currency_of_share_capital, date_of_last_annual_return,
                   financial_year_end_date, ultimate_holding_company_name
            FROM companies WHERE company_name_english ILIKE %s LIMIT 1
        """, (f"%{company_name}%",))
        row = cur.fetchone()
        cur.close(); conn.close()
        if row:
            return {
                "company_name_english": row[0], "company_registration_number": row[1],
                "registered_office_address": row[2], "principal_place_of_business": row[3],
                "status": row[4], "company_type": row[5],
                "directors": row[6] if isinstance(row[6], list) else [],
                "members": row[7] if isinstance(row[7], list) else [],
                "total_shares_issued": row[8], "currency_of_share_capital": row[9],
                "date_of_last_annual_return": str(row[10]) if row[10] else None,
                "financial_year_end_date": str(row[11]) if row[11] else None,
                "ultimate_holding_company_name": row[12],
            }
    except Exception:
        pass
    return None


def find_replacement(placeholder: str, data: dict[str, Any], template_name: str = None, company_name: str = None) -> str | None:
    """Find replacement value using learned field_mapping (if available) or fallback to data lookup."""
    placeholder_norm = placeholder.lower().replace(" ", "_").replace("-", "_")

    # === LEARNED MAPPING (Priority 1) ===
    if template_name:
        mapping = _get_field_mapping(template_name)
        if mapping:
            # Try exact match and normalized match
            config = mapping.get(placeholder) or mapping.get(placeholder_norm)
            if config:
                source = config.get("source", "")
                db_column = config.get("db_column", "")
                default = config.get("default", "TBD")

                if source == "db" and db_column:
                    # Get value from company DB
                    company_row = _get_company_from_db(company_name) if company_name else None
                    if company_row:
                        val = _get_company_field(company_row, db_column)
                        if val:
                            return val

                elif source == "user_input":
                    # Check if user provided this value in data
                    user_val = data.get(placeholder) or data.get(placeholder_norm)
                    if user_val and str(user_val) != "TBD":
                        return str(user_val)
                    # Use default
                    if default == "today":
                        from datetime import datetime
                        try:
                            import zoneinfo, os
                            tz = zoneinfo.ZoneInfo(os.getenv("TZ", "Asia/Yangon"))
                            return datetime.now(tz).strftime("%Y-%m-%d")
                        except Exception:
                            return datetime.now().strftime("%Y-%m-%d")
                    return default if default else "TBD"

    # === DATA LOOKUP (Priority 2 — fallback) ===
    # Direct match in provided data
    direct = data.get(placeholder) or data.get(placeholder_norm)
    if direct and str(direct) != "TBD":
        return str(direct)

    # Fuzzy match in data keys
    for key, value in data.items():
        if not value:
            continue
        key_norm = str(key).lower().replace(" ", "_").replace("-", "_")
        if key_norm == placeholder_norm or placeholder_norm in key_norm or key_norm in placeholder_norm:
            return str(value)

    # === SMART DEFAULTS (Priority 3) ===
    if "date" in placeholder_norm and "financial" not in placeholder_norm:
        from datetime import datetime
        try:
            import zoneinfo, os
            tz = zoneinfo.ZoneInfo(os.getenv("TZ", "Asia/Yangon"))
            return datetime.now(tz).strftime("%Y-%m-%d")
        except Exception:
            return datetime.now().strftime("%Y-%m-%d")

    if placeholder_norm == "pronoun":
        return "they"
    if "location" in placeholder_norm or "venue" in placeholder_norm:
        return data.get("registered_office", data.get("registered_office_address", "TBD"))

    return None


def generate_preview(doc: Document) -> str:
    """Generate preview of filled document."""
    preview_lines = ["Document Preview:"]

    for paragraph in doc.paragraphs[:10]:
        text = paragraph.text.strip()
        if text and len(text) > 10:
            preview_lines.append(f"  • {text[:100]}")

    return "\n".join(preview_lines)
